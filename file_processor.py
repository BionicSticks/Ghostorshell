import io
import os
import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import streamlit as st

class FileProcessor:
    def __init__(self):
        """Initialize the file processor"""
        # Set tesseract path if needed (usually auto-detected on Linux)
        pass
    
    def extract_text(self, uploaded_file):
        """
        Extract text from uploaded file based on its type
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text content
        """
        file_type = uploaded_file.type.lower()
        file_name = uploaded_file.name.lower()
        
        try:
            if file_type == "text/plain" or file_name.endswith('.txt'):
                return self._extract_text_from_txt(uploaded_file)
            elif file_type == "application/pdf" or file_name.endswith('.pdf'):
                return self._extract_text_from_pdf(uploaded_file)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith('.docx'):
                return self._extract_text_from_docx(uploaded_file)
            elif file_type.startswith("image/") or any(file_name.endswith(ext) for ext in ['.jpg', '.jpeg', '.png']):
                return self._extract_text_from_image(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {uploaded_file.name}: {str(e)}")
    
    def _extract_text_from_txt(self, uploaded_file):
        """Extract text from text file"""
        try:
            # Try different encodings
            content = uploaded_file.read()
            
            # Try UTF-8 first
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                try:
                    return content.decode('latin-1')
                except UnicodeDecodeError:
                    # Last resort - ignore errors
                    return content.decode('utf-8', errors='ignore')
                    
        except Exception as e:
            raise Exception(f"Could not read text file: {e}")
    
    def _extract_text_from_pdf(self, uploaded_file):
        """Extract text from PDF file using multiple methods"""
        text_content = ""
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Method 1: Try pdfplumber first (better for complex layouts)
            try:
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                
                if text_content.strip():
                    return text_content.strip()
            except Exception as e:
                st.warning(f"pdfplumber failed: {e}, trying PyPDF2...")
            
            # Method 2: Fallback to PyPDF2
            uploaded_file.seek(0)
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    return text_content.strip()
                else:
                    raise Exception("No text could be extracted from PDF")
                    
            except Exception as e:
                raise Exception(f"Both PDF extraction methods failed: {e}")
                
        except Exception as e:
            raise Exception(f"PDF processing error: {e}")
    
    def _extract_text_from_docx(self, uploaded_file):
        """Extract text from Word document (.docx)"""
        try:
            # Read the document
            doc = Document(uploaded_file)
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise Exception("No text content found in Word document")
            
            return full_text.strip()
            
        except Exception as e:
            raise Exception(f"Word document processing failed: {e}")
    
    def _extract_text_from_image(self, uploaded_file):
        """Extract text from image using OCR"""
        try:
            # Open image
            image = Image.open(uploaded_file)
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR with enhanced settings
            custom_config = r'--oem 3 --psm 6'
            extracted_text = pytesseract.image_to_string(image, config=custom_config, lang='eng')
            
            if not extracted_text.strip():
                # Try with different PSM mode if first attempt fails
                custom_config = r'--oem 3 --psm 3'
                extracted_text = pytesseract.image_to_string(image, config=custom_config, lang='eng')
            
            if not extracted_text.strip():
                raise Exception("No readable text found in image. Please ensure the image contains clear, readable text.")
            
            return extracted_text.strip()
            
        except pytesseract.TesseractNotFoundError:
            raise Exception("OCR service not available. Please try with text or PDF files instead.")
        except Exception as e:
            raise Exception(f"Image text extraction failed: {e}")
    
    def validate_file(self, uploaded_file, max_size_mb=10):
        """
        Validate uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            max_size_mb: Maximum file size in MB
            
        Returns:
            bool: True if valid, raises exception if invalid
        """
        if uploaded_file is None:
            raise ValueError("No file uploaded")
        
        # Check file size
        file_size_mb = uploaded_file.size / (1024 * 1024)
        if file_size_mb > max_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB. Maximum size: {max_size_mb}MB")
        
        # Check file type
        allowed_types = [
            "text/plain",
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "image/jpeg",
            "image/jpg",
            "image/png"
        ]
        
        allowed_extensions = ['.txt', '.pdf', '.docx', '.jpg', '.jpeg', '.png']
        
        file_type_valid = uploaded_file.type.lower() in allowed_types
        extension_valid = any(uploaded_file.name.lower().endswith(ext) for ext in allowed_extensions)
        
        if not (file_type_valid or extension_valid):
            raise ValueError(f"Unsupported file type: {uploaded_file.type}")
        
        return True
    
    def get_file_info(self, uploaded_file):
        """
        Get information about uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            dict: File information
        """
        return {
            'name': uploaded_file.name,
            'type': uploaded_file.type,
            'size': uploaded_file.size,
            'size_mb': uploaded_file.size / (1024 * 1024)
        }
